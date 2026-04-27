import os
import re
import json
from typing import Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# PDF Imports
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Register Tamil-compatible fonts from Windows
try:
    pdfmetrics.registerFont(TTFont('Nirmala', 'C:/Windows/Fonts/Nirmala.ttf'))
    pdfmetrics.registerFont(TTFont('Nirmala-Bold', 'C:/Windows/Fonts/NirmalaB.ttf'))
    FONT_REGULAR = "Nirmala"
    FONT_BOLD = "Nirmala-Bold"
except Exception as e:
    print(f"Font Load Warning: {e}. Falling back to Helvetica.")
    FONT_REGULAR = "Helvetica"
    FONT_BOLD = "Helvetica-Bold"

OUTPUTS_DIR = os.path.join(os.path.dirname(__file__), "..", "outputs")


# ── Main export function ───────────────────────────────────────
def export_documents(ncg_json: Any, session_id: str) -> tuple:
    os.makedirs(OUTPUTS_DIR, exist_ok=True)

    # Robustness: if ncg_json is a string, try to parse it
    if isinstance(ncg_json, str):
        try:
            parsed = json.loads(ncg_json)
            if isinstance(parsed, dict):
                ncg_json = parsed
            else:
                ncg_json = {"session_title": "Session Notes", "content": str(parsed)}
        except:
            ncg_json = {"session_title": "Session Notes", "content": ncg_json}

    # Generate clean filename from session title
    raw_title = (
        ncg_json.get("session_title") or
        ncg_json.get("title") or
        f"Notes_{session_id}"
    )
    # Clean for Windows filename
    clean_name = re.sub(r'[^\w\s-]', '', str(raw_title))
    clean_name = re.sub(r'\s+', '_', clean_name.strip())
    filename = clean_name[:80] if clean_name else "Session_Notes"

    pdf_path  = os.path.join(OUTPUTS_DIR, f"{filename}.pdf")
    docx_path = os.path.join(OUTPUTS_DIR, f"{filename}.docx")
    txt_path  = os.path.join(OUTPUTS_DIR, f"{filename}.txt")

    _generate_pdf(ncg_json, pdf_path)
    _generate_docx(ncg_json, docx_path)
    _generate_txt(ncg_json, txt_path)
    return f"/outputs/{filename}.pdf", f"/outputs/{filename}.docx"


# ── PDF Generator ──────────────────────────────────────────────
def _generate_pdf(ncg_json: dict, path: str):
    NAVY  = colors.HexColor("#284b63")
    CHAR  = colors.HexColor("#353535")
    BLUE_LIGHT = colors.HexColor("#f0f4f7")

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", fontSize=18, textColor=NAVY, spaceAfter=12, leading=22, alignment=TA_CENTER, fontName=FONT_BOLD)
    h2_style    = ParagraphStyle("H2", fontSize=14, textColor=NAVY, spaceBefore=15, spaceAfter=8, fontName=FONT_BOLD)
    h3_style    = ParagraphStyle("H3", fontSize=12, textColor=NAVY, spaceBefore=10, spaceAfter=5, fontName=FONT_REGULAR)
    body_style  = ParagraphStyle("Body", fontSize=10, textColor=CHAR, leading=14, fontName=FONT_REGULAR)
    bullet_style= ParagraphStyle("Bullet", fontSize=10, textColor=CHAR, leading=14, leftIndent=20, firstLineIndent=0, spaceBefore=2, fontName=FONT_REGULAR)
    label_style = ParagraphStyle("Label", fontSize=10, textColor=CHAR, fontName=FONT_BOLD)

    story = []

    # 1. Title Header
    title = (ncg_json.get("session_title") or ncg_json.get("title") or "Scribely Session Notes").upper()
    story.append(Paragraph(title, title_style))

    # Define metadata keys to skip in the main loop
    metadata_keys = {"session_title", "title", "prepared_by", "status", "session_id", "session_category"}

    sec_num = 1
    for key, data in ncg_json.items():
        if key in metadata_keys or not _has_content(data):
            continue

        label = key.replace("_", " ").title()
        story.append(Paragraph(f"{sec_num}. {label}", h2_style))
        story.append(HRFlowable(width="100%", thickness=1, color=NAVY, spaceAfter=10))

        # Specialized rendering for common keys
        if "details" in key.lower() and isinstance(data, dict):
            table_data = []
            for k, v in data.items():
                if isinstance(v, list): v = ", ".join(v)
                table_data.append([Paragraph(f"<b>{k.replace('_', ' ').title()}</b>", body_style), Paragraph(str(v), body_style)])
            t = Table(table_data, colWidths=[4*cm, 12*cm])
            t.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ('BACKGROUND', (0,0), (0,-1), BLUE_LIGHT),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('LEFTPADDING', (0,0), (-1,-1), 8),
            ]))
            story.append(t)

        elif key == "topics_covered" and isinstance(data, list):
            for idx, topic in enumerate(data, 1):
                if isinstance(topic, dict):
                    t_name = topic.get("topic_name") or topic.get("title") or "Topic"
                    if ":" in t_name:
                        pre, post = t_name.split(":", 1)
                        story.append(Paragraph(f"<b>{sec_num}.{idx} {pre}:</b>{post}", h3_style))
                    else:
                        story.append(Paragraph(f"<b>{sec_num}.{idx} {t_name}</b>", h3_style))
                    # Render topic sub-fields
                    for tk, tv in topic.items():
                        if tk in ["topic_name", "title"] or not tv: continue
                        label_sub = tk.replace("_", " ").title()
                        if isinstance(tv, list):
                            story.append(Paragraph(f"<b>{label_sub}:</b>", body_style))
                            for item in tv: story.append(Paragraph(str(item), bullet_style, bulletText="-"))
                        else:
                            story.append(Paragraph(f"<b>{label_sub}:</b> {tv}", body_style))
                else:
                    t_str = str(topic)
                    if ":" in t_str:
                        pre, post = t_str.split(":", 1)
                        story.append(Paragraph(f"<b>{sec_num}.{idx} {pre}:</b>{post}", h3_style))
                    else:
                        story.append(Paragraph(f"<b>{sec_num}.{idx} {t_str}</b>", h3_style))
                story.append(Spacer(1, 10))

        elif key == "qa_section" and isinstance(data, list):
            for i, qa in enumerate(data, 1):
                if isinstance(qa, dict):
                    story.append(Paragraph(f"<b>Q{i}:</b> {qa.get('question', '...')}", body_style))
                    if qa.get('answer'):
                        story.append(Paragraph(f"<b>A:</b> {qa.get('answer')}", body_style))
                else:
                    story.append(Paragraph(str(qa), bullet_style, bulletText="-"))
                story.append(Spacer(1, 8))

        elif isinstance(data, list):
            for item in data:
                story.append(Paragraph(str(item), bullet_style, bulletText="-"))
        
        elif isinstance(data, dict):
            for k, v in data.items():
                story.append(Paragraph(f"<b>{k.replace('_', ' ').title()}:</b> {v}", body_style))
        
        else:
            story.append(Paragraph(str(data), body_style))

        sec_num += 1
        story.append(Spacer(1, 15))

    doc.build(story)


# ══ DOCX Generation ════════════════════════════════════════════
def _generate_docx(notes: dict, path: str):
    doc = Document()
    NAVY = RGBColor(40, 75, 99)

    # Margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.8)
        section.left_margin = section.right_margin = Inches(0.8)

    # Title
    title_text = notes.get("session_title", "Class Notes").upper()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY

    metadata_keys = {"session_title", "title", "prepared_by", "status", "session_id"}

    sec_num = 1
    for key, data in notes.items():
        if key in metadata_keys or not _has_content(data):
            continue

        label = key.replace("_", " ").title()
        _docx_heading(doc, f"{sec_num}. {label}", NAVY)

        if "details" in key.lower() and isinstance(data, dict):
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            for k, v in data.items():
                row = table.add_row().cells
                row[0].text = k.replace("_", " ").title()
                row[0].paragraphs[0].runs[0].bold = True
                row[1].text = ", ".join(v) if isinstance(v, list) else str(v)

        elif key == "topics_covered" and isinstance(data, list):
            for idx, topic in enumerate(data, 1):
                if isinstance(topic, dict):
                    t_name = topic.get("topic_name") or topic.get("title") or "Topic"
                    p = doc.add_paragraph()
                    
                    if ":" in t_name:
                        pre, post = t_name.split(":", 1)
                        p.add_run(f"{sec_num}.{idx} {pre}:").bold = True
                        p.add_run(post)
                    else:
                        p.add_run(f"{sec_num}.{idx} {t_name}").bold = True
                    
                    for tk, tv in topic.items():
                        if tk in ["topic_name", "title"] or not tv: continue
                        label_sub = tk.replace('_', ' ').title()
                        if isinstance(tv, list):
                            doc.add_paragraph(f"{label_sub}:", style='Body Text').bold = True
                            for item in tv: doc.add_paragraph(str(item), style='List Bullet')
                        else:
                            p2 = doc.add_paragraph()
                            p2.add_run(f"{label_sub}: ").bold = True
                            p2.add_run(str(tv))
                else:
                    t_str = str(topic)
                    p = doc.add_paragraph()
                    if ":" in t_str:
                        pre, post = t_str.split(":", 1)
                        p.add_run(f"{sec_num}.{idx} {pre}:").bold = True
                        p.add_run(post)
                    else:
                        p.add_run(f"{sec_num}.{idx} {t_str}").bold = True

        elif key == "qa_section" and isinstance(data, list):
            for i, qa in enumerate(data, 1):
                if isinstance(qa, dict):
                    p = doc.add_paragraph()
                    p.add_run(f"Q{i}: ").bold = True
                    p.add_run(str(qa.get("question", "...")))
                    p = doc.add_paragraph()
                    p.add_run("A: ").bold = True
                    p.add_run(str(qa.get("answer", "")))
                else:
                    doc.add_paragraph(str(qa), style='List Bullet')

        elif isinstance(data, list):
            for item in data:
                doc.add_paragraph(str(item), style='List Bullet')
        
        elif isinstance(data, dict):
            for k, v in data.items():
                p = doc.add_paragraph()
                p.add_run(f"{k.replace('_', ' ').title()}: ").bold = True
                p.add_run(str(v))
        
        sec_num += 1
        doc.add_paragraph()

    doc.save(path)


def _docx_heading(doc: Document, title: str, color):
    para = doc.add_paragraph()
    run  = para.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = color
    # Border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "284B63")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _generate_txt(notes: dict, path: str):
    lines = []
    lines.append("=" * 60)
    lines.append(notes.get("session_title", "Notes").upper())
    lines.append("ONLINE CLASS SESSION NOTES")
    lines.append("=" * 60 + "\n")

    metadata_keys = {"session_title", "title", "prepared_by", "status", "session_id"}

    sec_num = 1
    for key, data in notes.items():
        if key in metadata_keys or not _has_content(data): continue
        lines.append(f"{sec_num}. {key.replace('_', ' ').title()}")
        lines.append("-" * 30)
        
        if isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    for k, v in item.items(): lines.append(f"  {k.title()}: {v}")
                else:
                    lines.append(f" • {item}")
        elif isinstance(data, dict):
            for k, v in data.items(): lines.append(f" {k.title()}: {v}")
        else:
            lines.append(str(data))
        lines.append("")
        sec_num += 1

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def _has_content(value) -> bool:
    if value is None: return False
    if isinstance(value, list): return len(value) > 0 and any(_has_content(v) for v in value)
    if isinstance(value, str): return len(value.strip()) > 0
    if isinstance(value, dict): return any(_has_content(v) for v in value.values())
    return bool(value)